from transcribe import transcribe_audio
from structure_yandex import structure_text
from docx import Document
from datetime import datetime
import os
from dotenv import load_dotenv
load_dotenv()

# 1. Расшифровка
text = transcribe_audio("example.mp3")  # Функция из transcribe.py

# 2. Структурируем
prompt = f"""Распредели по разделам:
1. Жалобы
2. Анамнез
3. Диагноз
4. Назначения
5. Рекомендации

Текст:
{text}
"""

data = structure_text(prompt)  # 🔁 теперь 'data' определена

# 3. Извлекаем результат
structured_text = data.get("result", {}).get("alternatives", [{}])[0].get("message", {}).get("text", "Ошибка в ответе")

# 4. Сохраняем в файлы
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
txt_path = f"output_{timestamp}.txt"
docx_path = f"output_{timestamp}.docx"

with open(txt_path, "w", encoding="utf-8") as f:
    f.write(structured_text)

doc = Document()
doc.add_heading("Структурированный текст приёма", 0)
for line in structured_text.split("\n"):
    doc.add_paragraph(line)
doc.save(docx_path)

print(f"\n✅ Сохранено в файлы:\n- {txt_path}\n- {docx_path}")
