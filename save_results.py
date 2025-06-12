from docx import Document
from datetime import datetime

def save_to_files(structured_text):
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    txt_path = f"output_{timestamp}.txt"
    docx_path = f"output_{timestamp}.docx"

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(structured_text)

    doc = Document()
    doc.add_heading("Структурированный текст приёма", 0)
    for line in structured_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(docx_path)

    return txt_path, docx_path
